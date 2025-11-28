import fitz  # PyMuPDF
import json
import networkx as nx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import re

def extract_text_from_pdf(path):
    """Extracts text from a PDF file."""
    text = ""
    with fitz.open(path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def parse_graph_data(relationship_json):
    """Parses entity relationship JSON into NetworkX DiGraph with vis.js-compatible properties."""
    if isinstance(relationship_json, str):
        relationship_json = json.loads(relationship_json)

    G = nx.DiGraph()

    for entity in relationship_json.get("entities", []):
        G.add_node(entity["id"], label=entity["name"], group=entity["type"])

    for rel in relationship_json.get("relationships", []):
        tooltip_lines = [
            f"Verb: {rel.get('verb', '')}",
            f"Optionality: {rel.get('Optionality', '')}",
            f"Condition: {rel.get('Condition for Relationship to be Active', '')}",
            f"Property: {rel.get('Property of Object (part of condition)', '')}",
            f"Thresholds: {rel.get('Thresholds', '')}",
            f"Frequency: {rel.get('frequency', '')}"
        ]
        tooltip_text = "\n".join(tooltip_lines)

        G.add_edge(
            rel["subject_id"],
            rel["object_id"],
            label=rel.get("verb", ""),
            title=tooltip_text
        )

    return G

def compare_graphs(G_old, G_new):
    """Compares old and new graphs to identify changed edges and nodes."""
    changed_edges = []
    added_nodes = list(set(G_new.nodes) - set(G_old.nodes))
    removed_nodes = list(set(G_old.nodes) - set(G_new.nodes))

    for u, v in G_new.edges:
        if G_old.has_edge(u, v):
            old_data = G_old[u][v]
            new_data = G_new[u][v]
            if old_data.get("label") != new_data.get("label") or old_data.get("title") != new_data.get("title"):
                changed_edges.append((u, v))
        else:
            changed_edges.append((u, v))

    return changed_edges, added_nodes, removed_nodes

def markdown_to_docx(doc: Document, text: str):
    lines = text.split('\n')
    table_buffer = []
    inside_table = False

    for line in lines:
        stripped = line.strip()

        # Handle empty lines
        if not stripped:
            if inside_table and table_buffer:
                insert_table_from_markdown(doc, table_buffer)
                table_buffer = []
                inside_table = False
            doc.add_paragraph()
            continue

        # Detect start of markdown table
        if stripped.startswith("|") and "|" in stripped[1:]:
            inside_table = True
            table_buffer.append(stripped)
            continue

        # End of table block
        if inside_table:
            insert_table_from_markdown(doc, table_buffer)
            table_buffer = []
            inside_table = False

        # Headings
        if stripped.startswith("###"):
            doc.add_heading(stripped.lstrip("#").strip(), level=3)
        elif stripped.startswith("##"):
            doc.add_heading(stripped.lstrip("#").strip(), level=2)
        elif stripped.startswith("#"):
            doc.add_heading(stripped.lstrip("#").strip(), level=1)

        # Bulleted or numbered lists
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')

        # Bold text
        elif "**" in stripped:
            para = doc.add_paragraph()
            while "**" in stripped:
                before, bold, rest = stripped.split("**", 2)
                para.add_run(before)
                bold_run = para.add_run(bold)
                bold_run.bold = True
                stripped = rest
            para.add_run(stripped)

        # Normal paragraph
        else:
            doc.add_paragraph(stripped)

    # Final table flush
    if inside_table and table_buffer:
        insert_table_from_markdown(doc, table_buffer)


def insert_table_from_markdown(doc: Document, lines: list):
    """
    Converts markdown-style table (| Col1 | Col2 | ...) to a real Word table with bold headers.
    Skips separator row with dashes.
    """
    rows = [re.split(r'\s*\|\s*', line.strip('|')) for line in lines if line.strip()]

    if len(rows) < 2:
        return  # Not a valid table

    headers = rows[0]

    # Detect if second row is a markdown separator (contains only dashes or colons)
    if len(rows) > 1 and all(re.match(r'^[:\-]+$', cell.strip()) for cell in rows[1]):
        data_rows = rows[2:]  # skip the second row
    else:
        data_rows = rows[1:]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Set bold headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        clean_header = header.strip().replace('**', '')
        run = hdr_cells[i].paragraphs[0].add_run(clean_header)
        run.bold = True

    for row in data_rows:
        if len(row) != len(headers):  # Skip malformed rows
            continue
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = cell_text.strip().replace('**', '')

