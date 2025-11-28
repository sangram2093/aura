import os
import json
from flask import Flask, request, render_template, redirect, url_for, jsonify, send_file, session
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from werkzeug.utils import secure_filename
from anthropic_llm import get_summary_with_context, get_entity_relationship_with_context, get_kop_doc
from utils import extract_text_from_pdf, parse_graph_data, markdown_to_docx
from db_models import db, Regulation, Upload, Summary, EntityGraph
from docx import Document
from io import BytesIO
import boto3

# Initialize Flask app
app = Flask(__name__)
app.secret_key = 'super-secret-key'

# Configurations
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///regulations.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Upload folder and file validation
UPLOAD_FOLDER = 'uploaded_pdfs'
ALLOWED_EXTENSIONS = {'pdf'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Initialize DB
db.init_app(app)

regulations_list = [                
                "EMIR Refit",
                "MiFID II",
                "SFTR",
                "AWPR",
                "AUSTRAC",
                "LME"
                ]
@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
       username = request.form.get('username', 'Guest') 
       session['logged_in'] = True
       session['username'] = username  
       return redirect(url_for('home'))
    return render_template("login.html")

@app.route("/home")
def home():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    username = session.get("username", "Guest")
    return render_template("home.html", username=username)

@app.route("/index")
def start():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    return redirect(url_for("index"))

@app.route("/logout", methods=["POST"])
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.route("/index.html", methods=["GET", "POST"])
def index():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    regulations = Regulation.query.all()

    regulations_name = [i.name for i in regulations]
    if request.method == "POST":
        regulation_id = request.form['regulation']
        mode = request.form.get('upload_mode')

        if mode == "first_time":
            new_file = request.files.get('first_time_file')
            if not new_file or not allowed_file(new_file.filename):
                return "Invalid or missing new regulation PDF.", 400

            filename_new = secure_filename(new_file.filename)
            new_path = os.path.join(app.config['UPLOAD_FOLDER'], filename_new)
            new_file.save(new_path)

            upload = Upload(regulation_id=regulation_id, old_path=None, new_path=new_path)

        else:
            old_file = request.files.get('old_file')
            new_file = request.files.get('new_file')

            if not old_file or not allowed_file(old_file.filename) or \
               not new_file or not allowed_file(new_file.filename):
                return "Invalid or missing old/new regulation PDFs.", 400

            filename_old = secure_filename(old_file.filename)
            filename_new = secure_filename(new_file.filename)

            old_path = os.path.join(app.config['UPLOAD_FOLDER'], filename_old)
            new_path = os.path.join(app.config['UPLOAD_FOLDER'], filename_new)

            old_file.save(old_path)
            new_file.save(new_path)

            upload = Upload(regulation_id=regulation_id, old_path=old_path, new_path=new_path)

        db.session.add(upload)
        db.session.commit()
        process_upload(upload.id)
        return redirect(url_for("compare", upload_id=upload.id))

    username = session.get("username", "Guest")
    return render_template("index.html", regulations=regulations,regulations_name = regulations_name, username=username)

@app.route('/dashboard')
def dashboard():
    return render_template('home.html')

# Upload History route
# @app.route('/history')
# def upload_history():
#     return render_template('history.html')

# Documentation PDF route
@app.route('/documentation')
def documentation():
    # Replace with the actual path to your PDF file
    pdf_path = 'static/docs/documentation.pdf'
    return send_file(pdf_path, as_attachment=False)

def process_upload(upload_id):
    upload = Upload.query.get(upload_id)

    db.session.query(Summary).filter_by(upload_id=upload.id).delete()
    db.session.query(EntityGraph).filter_by(upload_id=upload.id).delete()

    if not upload.old_path:
        new_text = extract_text_from_pdf(upload.new_path)
        new_summary = get_summary_with_context(new_text)
        new_json = get_entity_relationship_with_context(new_summary)

        G_new = parse_graph_data(json.loads(new_json))
        graph_new_json = json.dumps({
            "nodes": [{"id": n, **G_new.nodes[n]} for n in G_new.nodes],
            "edges": [{"from": u, "to": v, **G_new[u][v]} for u, v in G_new.edges]
        })

        db.session.add(Summary(upload_id=upload.id, old_summary=None, new_summary=new_summary))
        db.session.add(EntityGraph(
            upload_id=upload.id,
            old_json=None,
            new_json=new_json,
            graph_old=None,
            graph_new=graph_new_json
        ))
    else:
        old_text = extract_text_from_pdf(upload.old_path)
        new_text = extract_text_from_pdf(upload.new_path)

        old_summary = get_summary_with_context(old_text)
        new_summary = get_summary_with_context(new_text, context=old_summary)

        old_json = get_entity_relationship_with_context(old_summary)
        new_json = get_entity_relationship_with_context(new_summary, context=old_json)

        G_old = parse_graph_data(json.loads(old_json))
        G_new = parse_graph_data(json.loads(new_json))

        graph_old_json = json.dumps({
            "nodes": [{"id": n, **G_old.nodes[n]} for n in G_old.nodes],
            "edges": [{"from": u, "to": v, **G_old[u][v]} for u, v in G_old.edges]
        })

        graph_new_json = json.dumps({
            "nodes": [{"id": n, **G_new.nodes[n]} for n in G_new.nodes],
            "edges": [{"from": u, "to": v, **G_new[u][v]} for u, v in G_new.edges]
        })

        db.session.add(Summary(upload_id=upload.id, old_summary=old_summary, new_summary=new_summary))
        db.session.add(EntityGraph(
            upload_id=upload.id,
            old_json=old_json,
            new_json=new_json,
            graph_old=graph_old_json,
            graph_new=graph_new_json
        ))

    db.session.commit()

@app.route("/compare/<int:upload_id>")
def compare(upload_id):
    return render_template("compare.html", upload_id=upload_id)

@app.route("/graph_data/<int:upload_id>/<version>")
def graph_data(upload_id, version):
    graph_entry = EntityGraph.query.filter_by(upload_id=upload_id).first()
    if not graph_entry:
        return jsonify({"nodes": [], "edges": []})
    if version == "old":
        return jsonify(json.loads(graph_entry.graph_old or '{}'))
    if version == "new":
        return jsonify(json.loads(graph_entry.graph_new or '{}'))
    return jsonify({"error": "Invalid version"}), 400

@app.route("/regenerate/<int:upload_id>", methods=["POST"])
def regenerate(upload_id):
    process_upload(upload_id)
    return redirect(url_for("compare", upload_id=upload_id))

@app.route("/approve/<int:upload_id>", methods=["POST"])
def approve(upload_id):
    print(upload_id)
    summary = Summary.query.filter_by(upload_id=upload_id).first()
    graph = EntityGraph.query.filter_by(upload_id=upload_id).first()
    if not summary or not graph:
        return "Data not found", 404

    if not summary.new_summary or not graph.new_json:
        return "New data missing. Please upload new regulation first.", 400

    kop_text = get_kop_doc(new_summary=summary.new_summary, new_json_str=graph.new_json)
    doc = Document()
    doc.add_heading("Key Operating Procedure (KOP)", 0)
    markdown_to_docx(doc, kop_text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"kop_upload_{upload_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/generate_brd/<int:upload_id>", methods=["POST"])
def generate_brd(upload_id):
    summary = Summary.query.filter_by(upload_id=upload_id).first()
    graph = EntityGraph.query.filter_by(upload_id=upload_id).first()
    if not summary or not graph:
        return "Data not found", 404

    if not summary.new_summary or not graph.new_json:
        return "New data missing. Please upload new regulation first.", 400

    # You will need to implement this function in anthropic_llm.py
    from anthropic_llm import get_brd_doc
    brd_text = get_brd_doc(new_summary=summary.new_summary, new_json_str=graph.new_json)

    doc = Document()
    doc.add_heading("Business Requirements Document (BRD)", 0)
    markdown_to_docx(doc, brd_text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"brd_upload_{upload_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/history")
def history():
    uploads = Upload.query.order_by(Upload.upload_time.desc()).all()
    return render_template("history.html", uploads=uploads)



#new code for chatbot
BEDROCK_REGION = os.environ.get("AWS_REGION", "us-east-2")
KB_ID = os.environ.get("BEDROCK_KB_ID", "MAKZOATKHX")
MODEL_ARN = os.environ.get("MODEL_ARN",
                           "arn:aws:bedrock:us-east-2::foundation-model/anthropic.claude-3-haiku-20240307-v1:0")

# Create Bedrock Agent Runtime client
bedrock_agent_runtime = boto3.client("bedrock-agent-runtime", region_name=BEDROCK_REGION)


@app.route("/ask", methods=["POST"])
def ask_question():
    data = request.get_json()
    question = data.get("question")

    if not question:
        return jsonify({"error": "Missing question"}), 400

    try:
        response = bedrock_agent_runtime.retrieve_and_generate(
            input={"text": question},
            retrieveAndGenerateConfiguration={
                "type": "KNOWLEDGE_BASE",
                "knowledgeBaseConfiguration": {
                    "knowledgeBaseId": KB_ID,
                    "modelArn": "arn:aws:bedrock:us-east-2:908924925940:inference-profile/us.anthropic.claude-3-haiku-20240307-v1:0"
                }
            }
        )
        answer = response['output']['text']
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        if not Regulation.query.first():
            regulations_objs = []
            for regulation_name in regulations_list:
                print(regulation_name)
                regulations_objs.append(Regulation(name=regulation_name))
            db.session.add_all(regulations_objs)
            db.session.commit()

    app.run(debug=True, host="0.0.0.0", port=8080)